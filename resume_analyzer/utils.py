import re
import pdfplumber
import PyPDF2
import docx
from io import BytesIO

def extract_text_from_pdf(file_file):
    text = ""
    try:
        # Try pdfplumber first
        with pdfplumber.open(file_file) as pdf:
            for page in pdf.pages:
                text_page = page.extract_text()
                if text_page:
                    text += text_page + "\n"
    except Exception:
        # Fallback to PyPDF2
        try:
            file_file.seek(0)
            pdf_reader = PyPDF2.PdfReader(file_file)
            for page in pdf_reader.pages:
                text_page = page.extract_text()
                if text_page:
                    text += text_page + "\n"
        except Exception as e:
            text = f"Error reading PDF: {str(e)}"
    return text

def extract_text_from_docx(file_file):
    text = ""
    try:
        doc = docx.Document(file_file)
        for para in doc.paragraphs:
            text += para.text + "\n"
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text += cell.text + " "
                text += "\n"
    except Exception as e:
        text = f"Error reading DOCX: {str(e)}"
    return text

def parse_resume_text(text, master_skills, api_key=None):
    if api_key:
        try:
            import urllib.request
            import json
            import time

            prompt = (
                f"You are an expert resume parser. Analyze the following resume text:\n\n"
                f"{text}\n\n"
                f"Extract the following information and return it strictly as a JSON object with these keys:\n"
                f"- 'name': The candidate's name (string)\n"
                f"- 'email': The candidate's email address (string)\n"
                f"- 'phone': The candidate's phone number (string)\n"
                f"- 'education': A summary of candidate's education history (string, bulleted or formatted)\n"
                f"- 'certifications': A summary of candidate's certifications (string, bulleted or formatted)\n"
                f"- 'skills': A list of professional skills, domain competencies, tools, or methodologies found in the resume (list of strings). For example: AutoCAD, Surveying, RCC Design for Civil Engineering; Financial Modeling, Marketing Strategy, Operations Management for MBA; Contract Drafting, Legal Research, Corporate Law for Law; Auditing, Taxation, Tally, Bookkeeping for B.Com; Python, Django, Docker for IT, etc. Extract all relevant professional skills and do not limit yourself.\n"
                f"- 'primary_domain': The main professional field/domain of the candidate based on their skills and experience (e.g., 'Civil Engineering', 'DevOps Engineering', 'Backend Development', 'Business Administration', 'Legal Practice', 'Accounting & Finance', etc.) (string)\n"
                f"- 'career_goal': A suitable career goal/role for this candidate (e.g., 'Civil Engineer', 'DevOps Engineer', 'Backend Developer', 'Management Professional', 'Legal Counsel', 'Financial Analyst', etc.) (string)\n\n"
                f"Return ONLY the raw JSON object, without any markdown formatting or backticks."
            )
            data = {
                "contents": [{
                    "parts": [{
                        "text": prompt
                    }]
                }]
            }

            models = [
                "gemini-3.5-flash",
                "gemini-3.1-flash-lite",
                "gemini-2.5-flash",
                "gemini-2.0-flash",
                "gemini-2.0-flash-lite",
                "gemini-1.5-flash",
                "gemini-flash-latest"
            ]

            parsed_json = None
            last_error = None

            for model in models:
                for api_version in ["v1beta", "v1"]:
                    url = f"https://generativelanguage.googleapis.com/{api_version}/models/{model}:generateContent?key={api_key}"
                    for attempt in range(2):
                        try:
                            req = urllib.request.Request(
                                url,
                                data=json.dumps(data).encode('utf-8'),
                                headers={
                                    'Content-Type': 'application/json',
                                    'x-goog-api-key': api_key
                                },
                                method='POST'
                            )
                            with urllib.request.urlopen(req, timeout=15) as response:
                                res_data = json.loads(response.read().decode('utf-8'))
                                text_response = res_data['candidates'][0]['content']['parts'][0]['text'].strip()
                                
                                # Clean possible markdown block markers
                                if text_response.startswith("```"):
                                    lines = text_response.splitlines()
                                    if lines[0].startswith("```"):
                                        lines = lines[1:]
                                    if lines and lines[-1].startswith("```"):
                                        lines = lines[:-1]
                                    text_response = "\n".join(lines).strip()
                                
                                parsed_json = json.loads(text_response)
                                break
                        except Exception as e:
                            last_error = e
                            error_msg = str(e)
                            if "404" in error_msg or "400" in error_msg or "403" in error_msg or "429" in error_msg:
                                break
                            time.sleep(1)
                    if parsed_json is not None:
                        break
                if parsed_json is not None:
                    break

            if parsed_json is not None:
                def clean_str_field(val):
                    if val is None:
                        return ""
                    if isinstance(val, list):
                        return ", ".join(clean_str_field(item) for item in val if item is not None)
                    if isinstance(val, dict):
                        return ", ".join(f"{k}: {clean_str_field(v)}" for k, v in val.items() if v is not None)
                    return str(val).strip()

                skills_list = parsed_json.get('skills', [])
                cleaned_skills = []
                if isinstance(skills_list, list):
                    for s in skills_list:
                        if isinstance(s, str):
                            cleaned_skills.append(s.strip())
                        elif isinstance(s, list):
                            for sub_s in s:
                                if isinstance(sub_s, str):
                                    cleaned_skills.append(sub_s.strip())
                        elif isinstance(s, dict):
                            for val in s.values():
                                if isinstance(val, str):
                                    cleaned_skills.append(val.strip())
                elif isinstance(skills_list, str):
                    cleaned_skills = [s.strip() for s in skills_list.split(',') if s.strip()]
                
                if not cleaned_skills:
                    text_lower = text.lower()
                    for skill in master_skills:
                        if re.search(r'\b' + re.escape(skill.lower()) + r'\b', text_lower):
                            cleaned_skills.append(skill)

                return {
                    'name': clean_str_field(parsed_json.get('name')),
                    'email': clean_str_field(parsed_json.get('email')),
                    'phone': clean_str_field(parsed_json.get('phone')),
                    'education': clean_str_field(parsed_json.get('education')) or "Not clearly identified",
                    'certifications': clean_str_field(parsed_json.get('certifications')) or "Not clearly identified",
                    'skills': cleaned_skills,
                    'primary_domain': clean_str_field(parsed_json.get('primary_domain')) or "Software Engineering",
                    'career_goal': clean_str_field(parsed_json.get('career_goal')) or "Software Engineer",
                    'extracted_by': 'AI'
                }
            else:
                print(f"Gemini Resume Parsing failed: {str(last_error)}. Falling back to local parser.")
        except Exception as e:
            print(f"Gemini Resume Parsing error during setup: {str(e)}. Falling back to local parser.")

    parsed_data = {
        'name': None,
        'email': None,
        'phone': None,
        'education': [],
        'certifications': [],
        'skills': [],
        'primary_domain': 'Software Engineering',
        'career_goal': 'Software Engineer',
        'extracted_by': 'Code'
    }

    # Extract Email
    email_match = re.search(r'[\w\.-]+@[\w\.-]+\.\w+', text)
    if email_match:
        parsed_data['email'] = email_match.group(0)

    # Extract Phone (common patterns including international and spaces)
    phone_match = re.search(r'\+?\d[\d\s()-]{8,14}\d', text)
    if phone_match:
        parsed_data['phone'] = phone_match.group(0).strip()

    # Extract Name Heuristic
    # Look at the first 5 non-empty lines, skip lines with email/phone/websites, choose the first clean line
    lines = [line.strip() for line in text.split('\n') if line.strip()]
    for line in lines[:5]:
        if parsed_data['email'] and parsed_data['email'] in line:
            continue
        if parsed_data['phone'] and parsed_data['phone'] in line:
            continue
        if any(keyword in line.lower() for keyword in ['curriculum', 'resume', 'cv', 'profile', 'http', 'www', 'portfolio']):
            continue
        # Name should be relatively short
        if 3 <= len(line.split()) <= 4 and re.match(r'^[a-zA-Z\s]+$', line):
            parsed_data['name'] = line
            break
    
    # Fallback to first line if no clean name found
    if not parsed_data['name'] and lines:
        parsed_data['name'] = lines[0][:100]

    # Extract Skills
    text_lower = text.lower()
    for skill in master_skills:
        # Use word boundaries to prevent substring matches (like 'Go' matching in 'Google')
        pattern = r'\b' + re.escape(skill.lower()) + r'\b'
        if re.search(pattern, text_lower):
            parsed_data['skills'].append(skill)

    # Section extraction heuristics for Education & Certifications
    current_section = None
    education_keywords = ['education', 'academic', 'qualification', 'degree', 'schooling', 'university', 'college']
    cert_keywords = ['certification', 'certificates', 'credential', 'certified', 'licensure']

    sections = {
        'education': [],
        'certifications': []
    }

    # Line by line processing to capture sections
    for line in lines:
        line_lower = line.lower()
        
        # Check if line indicates a section change
        if any(kw in line_lower for kw in education_keywords) and len(line) < 30:
            current_section = 'education'
            continue
        elif any(kw in line_lower for kw in cert_keywords) and len(line) < 30:
            current_section = 'certifications'
            continue
        elif len(line) < 30 and any(keyword in line_lower for keyword in ['experience', 'project', 'skills', 'objective', 'summary', 'contact']):
            current_section = None
            continue
        
        # Add content to the active section
        if current_section:
            sections[current_section].append(line)

    parsed_data['education'] = "\n".join(sections['education'][:6]) if sections['education'] else "Not clearly identified"
    parsed_data['certifications'] = "\n".join(sections['certifications'][:6]) if sections['certifications'] else "Not clearly identified"

    # Detect domain/goal from skills and text for local fallback parser
    skills_lower = [s.lower() for s in parsed_data['skills']]
    text_lower = text.lower()
    
    civil_keywords = ['autocad', 'staad', 'etabs', 'surveying', 'estimation', 'rcc', 'quantity surveying', 'site supervision', 'primavera']
    devops_keywords = ['docker', 'aws', 'jenkins', 'kubernetes', 'terraform', 'ansible', 'ci/cd']
    ml_keywords = ['machine learning', 'data science', 'nltk', 'scikit-learn', 'tensorflow', 'pytorch', 'pandas', 'numpy']
    mba_keywords = ['marketing', 'finance', 'operations', 'strategy', 'human resources', 'management', 'business development', 'sales', 'mba']
    law_keywords = ['law', 'legal', 'contract', 'drafting', 'constitution', 'advocate', 'litigation', 'court', 'corporate law', 'jurisprudence']
    bcom_keywords = ['accounting', 'audit', 'taxation', 'tally', 'finance', 'gst', 'ledger', 'balance sheet', 'bookkeeping', 'bcom']

    is_civil = any(any(kw in s for kw in civil_keywords) for s in skills_lower) or any(kw in text_lower for kw in civil_keywords)
    is_devops = any(any(kw in s for kw in devops_keywords) for s in skills_lower) or any(kw in text_lower for kw in devops_keywords)
    is_ml = any(any(kw in s for kw in ml_keywords) for s in skills_lower) or any(kw in text_lower for kw in ml_keywords)
    is_mba = any(any(kw in s for kw in mba_keywords) for s in skills_lower) or any(kw in text_lower for kw in mba_keywords)
    is_law = any(any(kw in s for kw in law_keywords) for s in skills_lower) or any(kw in text_lower for kw in law_keywords)
    is_bcom = any(any(kw in s for kw in bcom_keywords) for s in skills_lower) or any(kw in text_lower for kw in bcom_keywords)

    if is_civil:
        parsed_data['primary_domain'] = 'Civil Engineering'
        parsed_data['career_goal'] = 'Civil Engineer'
    elif is_devops:
        parsed_data['primary_domain'] = 'DevOps Engineering'
        parsed_data['career_goal'] = 'DevOps Engineer'
    elif is_ml:
        parsed_data['primary_domain'] = 'Data Science'
        parsed_data['career_goal'] = 'Data Scientist'
    elif is_mba:
        parsed_data['primary_domain'] = 'Business Administration'
        parsed_data['career_goal'] = 'Management Professional'
    elif is_law:
        parsed_data['primary_domain'] = 'Legal Practice'
        parsed_data['career_goal'] = 'Legal Counsel'
    elif is_bcom:
        parsed_data['primary_domain'] = 'Accounting & Finance'
        parsed_data['career_goal'] = 'Financial Analyst'
    elif any(kw in ''.join(skills_lower) or kw in text_lower for kw in ['python', 'django', 'mysql', 'sql', 'software']):
        parsed_data['primary_domain'] = 'Backend Development'
        parsed_data['career_goal'] = 'Backend Developer'

    return parsed_data
